import os
import sys
from os import remove
from posixpath import dirname
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.select import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC, wait
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Cm, Inches
from tkinter import *
from PIL import ImageTk, Image
import time
import re

def login():
    #Login Page
    username = driver.find_element_by_xpath('//*[@id="login_user_name"]')
    password = driver.find_element_by_xpath('//*[@id="login_pwd"]')
    login = driver.find_element_by_xpath('//*[@id="doLogin"]')
    username.send_keys("")
    password.send_keys("")
    login.click()
    
def goTofullfillment():
    #Go to fullfillment page
    driver.find_element_by_xpath('//*[@id="menu-item-19236"]/a').click()
    driver.find_element_by_xpath('//*[@id="navBarToggler"]/ul[1]/li[2]/a').click()
    driver.find_element_by_xpath('//*[@id="navBarToggler"]/ul[1]/li[2]/ul/li[2]/a').click()

#Repeated click until on element until it exists (helps with slow loading times)
def tryClick(xpath):
    while True:
        try:
            driver.find_element_by_xpath(xpath).click()
            break
        except:
            time.sleep(.1)

#Filter list by route
def filterByRoute(routeName):
    tryClick('//*[@id="orders_grid"]/table/thead/tr/th[2]/a[1]/span')
    waitTime = .1
    while True:
        try:
            tryClick('//*[(@class="k-link") and (contains(text(), "Filter"))]')
            filterInput = driver.find_element_by_xpath('//*[(@class="k-textbox") and (@title="Value")]')
            time.sleep(waitTime)
            filterInput.clear()
            filterInput.send_keys(routeName)
            time.sleep(waitTime)
            submit = driver.find_element_by_xpath('//*[(@type="submit") and (@class="k-button k-primary")]')
            submit.click()
            break
        except:
            waitTime += .1

#Gets picklist for given route
def getRouteInv(routeName, routeSales, routeFreebies, freebiePrice):
    filterByRoute(routeName)
    numOrders = getNumOrders()
    for i in range(numOrders):
        time.sleep(.3)
        itemNum = 1
        while hasXpath('//*[@id="orders_items_grid"]/table/tbody/tr[{0}]/td[2]'.format(itemNum)):
            curItem = driver.find_element_by_xpath('//*[@id="orders_items_grid"]/table/tbody/tr[{0}]/td[2]'.format(itemNum)).text
            itemQuantity = int(driver.find_element_by_xpath('//*[@id="orders_items_grid"]/table/tbody/tr[{0}]/td[4]'.format(itemNum)).text)
            addItem(curItem, itemQuantity)
            itemNum += 1
        time.sleep(.1)
        priceElem = driver.find_element_by_xpath('//*[@id="orders_items_grid"]/table/tfoot/tr/td[5]/div')
        priceStr = priceElem.text
        priceStr = priceStr[1:]
        priceFloat = float(priceStr)
        if priceFloat >= freebiePrice:
            routeFreebies += 1
        routeSales += priceFloat
        driver.find_element_by_xpath('//*[@id="orderDetailsPanel"]/div[1]/div/div[2]/div[2]/a[2]').click()
    return routeSales, routeFreebies


#Gets number of orders for current filter
def getNumOrders():
    numOrderElem = driver.find_element_by_xpath('//*[@id="orders_grid"]/div[1]/span[2]')
    numOrderStr = numOrderElem.text
    tempNumOrder = re.findall(r'\d+', numOrderStr)
    numOrderList = list(map(int, tempNumOrder))
    numOrderInt = numOrderList[2]
    return numOrderInt

def getNumRoutes():
    while not hasXpath('//*[@id="grid"]/div[1]/span[2]'):
        pass
    numRoutesElem = driver.find_element_by_xpath('//*[@id="grid"]/div[1]/span[2]')
    numRoutesStr = numRoutesElem.text
    tempNumRoutes = re.findall(r'\d+', numRoutesStr)
    numRoutesList = list(map(int, tempNumRoutes))
    numRoutesInt = numRoutesList[2]
    return numRoutesInt

#Adds item to picklist
def addItem(itemName, quantity):
    if itemName in pickList:
        pickList[itemName] += quantity
    else:
        pickList[itemName] = quantity

#Checks if xpath exists
def hasXpath(xpath):
    try:
        driver.find_element_by_xpath(xpath)
        return True
    except:
        return False

def resize_image(e):
    global image, resized, image2
    # open image to resize it
    image = Image.open(os.path.join(application_path, "picklistBG.jpg"))
    # resize the image with width and height of root
    resized = image.resize((e.width, e.height), Image.ANTIALIAS)

    image2 = ImageTk.PhotoImage(resized)
    canvas.create_image(0, 0, image=image2, anchor='nw')

#Starts webscraping after confirm button is clicked
def confirmClick():
    global driver
    root.destroy()
    driver = webdriver.Chrome()
    driver.maximize_window()
    totalSales = 0.0
    totalFreebies = 0

    #Get selected routes from user
    for route in dropVars:
        selectedRoutes.append(route.get())

    #Login to rep portal
    driver.get("https://repportal.collectivegoods.com/")
    login()
    goTofullfillment()

    #Get total sales (including all routes)
    #This helps to determine if new sales come in later
    tryClick('//*[@id="orders_grid"]/table/thead/tr/th[2]/a[1]/span') #Used to discover when page has loaded, so unable to locate element error doesnt occur
    tryClick('/html/body/section/div/div/div/div[3]/div[1]/div/div[1]/h3[1]')
    totalOrders = getNumOrders()
    time.sleep(1)

    #Get inventory for each route in selectedRoutes
    for route in selectedRoutes:
        routeSales = 0
        routeFreebies = 0
        routeSales, routeFreebies = getRouteInv(route, routeSales, routeFreebies, float(freebiePrice.get()))
        totalSales += routeSales
        totalFreebies += routeFreebies


    driver.close()

    #Create inventory document and save in same directory as this application
    document = Document()
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Qty"
    for item in pickList:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[0].width = Inches(4)
        row_cells[1].text = str(pickList[item])
    document.add_paragraph('Total Sales: $' + str(totalSales))
    document.add_paragraph('Total Freebies: ' + str(totalFreebies))
    document.add_paragraph('Total Orders (includes all routes): ' + str(totalOrders))

    #Name of file is current date + route names
    routesStr = '-'.join([str(elem) for elem in selectedRoutes])
    routesStr = routesStr.replace("/", "-")
    document.save(time.strftime(os.path.join(application_path, "%Y-%m-%d " + routesStr + ".docx")))

#Displays dropdown menus for the number of routes selected
def genDrops(event):
    #Destroy all current dropdown menus and reset dropdown variables array
    for i in routesMenu:
        i.destroy()
    dropVars.clear()

    #Repopulate array with variables for each dropdown menu
    for i in range(numRoutes.get()):
        dropVars.append(StringVar())

    #Recreate all dropdown menus for route selection
    for i in range(numRoutes.get()):
        curDrop = OptionMenu(root, dropVars[i], *routes)
        dropCanvas = canvas.create_window(0, 40*i, anchor = "nw", window = curDrop)
        routesMenu.append(curDrop)
        canvasArr.append(dropCanvas)

def getRoutes():
    root.destroy()
    global driver
    driver = webdriver.Chrome()
    driver.maximize_window()
    driver.get("https://repportal.collectivegoods.com/")
    login()

    #Go to routes page
    driver.find_element_by_xpath('//*[@id="menu-item-19236"]/a').click()
    driver.find_element_by_xpath('//*[@id="navBarToggler"]/ul[1]/li[1]/a').click()
    driver.find_element_by_xpath('//*[@id="navBarToggler"]/ul[1]/li[1]/ul/li[1]/a').click()

    f = open(os.path.join(application_path, "All Routes.txt"), "w")
    totalRoutes = getNumRoutes()
    for i in range(totalRoutes):
        if i % 20 == 0 and i != 0:
            driver.find_element_by_xpath('//*[@id="grid"]/div[1]/a[3]/span').click()
        curRoute = driver.find_element_by_xpath('//*[@id="grid"]/table/tbody/tr[{0}]/td[1]'.format((i % 20) + 1)).text
        f.write(curRoute + "\n")
        time.sleep(.1)
    f.close()
    driver.close()

    #f = open(os.path.join(application_path, "All Routes.txt"))

if __name__ == "__main__":
    #Dictionary for products and their quantity
    pickList = {}
    
    #Determine if file is running from executable or python script
    if getattr(sys, 'frozen', False):
        application_path = os.path.dirname(sys.executable)
    elif __file__:
        application_path = os.path.dirname(__file__)

    #Initialize application
    root = Tk()
    root.title("CG - Picklist")
    root.iconbitmap(os.path.join(application_path, "cglogo.ico"))

    #Get screen dimensions
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    minDim = min(screen_height, screen_width)
    maxDim = max(screen_height, screen_width)

    #Center application on screen
    if(minDim == screen_height):
        #Landscape monitor
        root.geometry('{0}x{1}+{2}+{3}'.format(300, 300, int((maxDim/2)-150), int((minDim/2)-150)))
    else:
        #Portrait monitor
        root.geometry('{0}x{1}+{2}+{3}'.format(300, 300, int((minDim/2)-150)), int((maxDim/2)-150))

    #Adjust image size to fit application window
    canvas = Canvas(root, width = minDim, height = minDim)
    canvas.pack(fill = "both", expand = True)
    image = Image.open(os.path.join(application_path, "picklistBG.jpg"))
    resized = image.resize((300, 300), Image.ANTIALIAS)
    bg = ImageTk.PhotoImage(resized)
    canvas.create_image(0,0, image = bg, anchor = "nw")

    selectedRoutes = []
    routes = []
    try:
        f = open(os.path.join(application_path, "All Routes.txt"), "r")
        for route in f:
            routes.append(route[:-1])
    except:
        f = open(os.path.join(application_path, "All Routes.txt"), "w")
        f.write("None\n")
        f.close()
        routes.append("None")

    #Update routes
    getRoutesButton = Button(root, text="Update Routes", command=getRoutes, bg="#FFFFFF", font="Helvetica 10 bold")
    routesButtonCanvas = canvas.create_window(245, 20, anchor = "center", window = getRoutesButton)

    #Determine number of routes
    numOptions = [1, 2, 3, 4, 5]
    numRoutes = IntVar()
    numRoutes.set(numOptions[0])
    numRoutesDrop = OptionMenu(root, numRoutes, *numOptions, command=genDrops)
    numRoutesDropCanvas = canvas.create_window(135, 260, anchor = "nw", window = numRoutesDrop)

    #Determine Freebie Price
    freebiePrice = StringVar()
    freebieEntry = Entry(root, textvariable=freebiePrice, width=7)
    freebieEntryCanvas = canvas.create_window(135, 240, anchor="nw", window = freebieEntry)

    #Routes Label
    canvas.create_text(5, 240, text="Freebie Price Point:", anchor="nw", font="Helvetica 10 bold", fill="white")
    canvas.create_text(5, 270, text="Number of Routes:", anchor="nw", font="Helvetica 10 bold", fill="white")

    #Corresponding dropdown values
    dropVars = []
    routesMenu = []
    canvasArr = []

    #Confirm button to begin execution
    confirmButton = Button(root, text = "Confirm", command=confirmClick, bg="#FFFFFF", font="Helvetica 12 bold")
    confirmButtonCanvas = canvas.create_window(210, 260, anchor = "nw", window = confirmButton)

    root.mainloop()